"""Word export for the overall study plan.

Mirrors the production layout: title + 2nd major, grouped by year (one section
per year), each semester/summer shown as a Course/Credits table with a subtotal,
and a year-total at the bottom. Adapted to the refactored unified course table
(`Course Credits Study Period`) instead of the old CUHK/CUHKSZ split columns.
"""

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import pandas as pd

from app.constant import STUDY_CAMPUS


def generate_study_plan_docx(study_plan_df: pd.DataFrame, major_2_name: str) -> BytesIO:
    """Generate a Word document from the unified study-plan DataFrame.

    Expected columns: Course, Credits, Study Period.
    """
    doc = Document()

    # Title
    title = doc.add_heading("IDADM Study Plan", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Second Major Info
    p = doc.add_paragraph()
    p.add_run("Second Major: ").bold = True
    p.add_run(major_2_name)

    # Generation Date
    p = doc.add_paragraph()
    p.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for year in range(1, 5):
        doc.add_heading(f"Year {year}", level=1)

        periods = [f"Year {year} Sem 1", f"Year {year} Sem 2"]
        if year < 4:
            periods.extend([f"Year {year} Summer (CUHK)", f"Year {year} Summer (CUHKSZ)"])

        year_total_credits = 0

        for period in periods:
            period_df = study_plan_df[study_plan_df["Study Period"] == period]

            if not period_df.empty:
                # Semester heading
                p = doc.add_paragraph()
                title = f"{period} ({STUDY_CAMPUS.get(period, 'N/A')})" if STUDY_CAMPUS.get(period, "") not in period else period
                run = p.add_run(title)
                run.bold = True
                run.font.size = Pt(12)

                # Table
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                table.autofit = False

                table.columns[0].width = Inches(5.0)
                table.columns[1].width = Inches(1.0)

                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Course"
                hdr_cells[0].width = Inches(5.0)
                hdr_cells[1].text = "Credits"
                hdr_cells[1].width = Inches(1.0)
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Data rows
                display_df = period_df.filter(["Course", "Credits"])
                for _, data_row in display_df.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(data_row["Course"])
                    row_cells[0].width = Inches(5.0)
                    row_cells[1].text = str(data_row["Credits"])
                    row_cells[1].width = Inches(1.0)
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Subtotal
                period_credits = int(period_df["Credits"].sum())
                year_total_credits += period_credits
                p = doc.add_paragraph()
                p.add_run(f"Subtotal Credits: {period_credits}").italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph()  # Spacer

        # Year Total
        p = doc.add_paragraph()
        run = p.add_run(f"Year {year} Total Credits: {year_total_credits}")
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if year < 4:
            doc.add_page_break()

    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream
